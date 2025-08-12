#!/usr/bin/env python3
"""
generate_assessment.py

Generates:
 - two images (uniform table, packed-spheres top view),
 - a Word (.docx) file containing two MCQs in the given Question Output Format,
 - a github-ready folder with QUESTIONS.md, README.md and images,
 - a zip file of that folder.

Optionally initializes a git repo and pushes to a remote URL provided via --push-repo.
"""

import argparse
import os
import shutil
import zipfile
import subprocess
from pathlib import Path

# Image libs
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt

# Document libs
from docx import Document
from docx.shared import Inches

def make_uniform_image(path: Path):
    """Create a simple visual table image listing shirt/pants/hat colors."""
    img = Image.new("RGB", (900, 280), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None

    draw.text((10, 10), "Available Options", fill="black", font=font)
    draw.text((20, 40), "Shirts:", fill="black", font=font)
    shirts = ["Tan", "Red", "White", "Yellow"]
    pants = ["Black", "Khaki", "Navy"]
    hats = ["Blue", "Green", "Brown"]
    y = 70
    box_w, box_h = 200, 30
    for s in shirts:
        draw.rectangle([20, y, 20 + box_w, y + box_h], outline="black")
        draw.text((25, y + 6), s, fill="black", font=font)
        y += 40

    y = 70
    for p in pants:
        draw.rectangle([260, y, 260 + box_w, y + box_h], outline="black")
        draw.text((265, y + 6), p, fill="black", font=font)
        y += 40

    y = 70
    for h in hats:
        draw.rectangle([500, y, 500 + box_w, y + box_h], outline="black")
        draw.text((505, y + 6), h, fill="black", font=font)
        y += 40

    img.save(path)
    print(f"Created {path}")

def make_packed_spheres_image(path: Path, rows=2, cols=4, radius_cm=2):
    """
    Create a top-view image of tightly packed circles arranged in grid rows x cols.
    radius_cm is used for dimension labeling (visual).
    """
    scale = 30  # pixels per cm approx
    diameter = 2 * radius_cm
    r_px = radius_cm * scale
    width_px = int(cols * diameter * scale)
    height_px = int(rows * diameter * scale)

    fig, ax = plt.subplots(figsize=(width_px/100, height_px/100), dpi=100)
    for i in range(rows):
        for j in range(cols):
            cx = (j * diameter * scale) + r_px
            cy = (i * diameter * scale) + r_px
            circle = plt.Circle((cx, cy), r_px, edgecolor='black', fill=False, linewidth=1.5)
            ax.add_patch(circle)

    ax.set_xlim(0, width_px)
    ax.set_ylim(0, height_px)
    ax.set_aspect('equal')
    ax.axis('off')
    plt.tight_layout(pad=0)
    plt.savefig(path, bbox_inches='tight', pad_inches=0.02)
    plt.close()
    print(f"Created {path}")

def create_docx(docx_path: Path, img1: Path, img2: Path):
    doc = Document()
    doc.add_heading('Generated Assessment Questions', level=1)
    doc.add_paragraph('@title Central Middle — Derived Assessment')
    doc.add_paragraph('@description Two generated quantitative math questions similar to the base examples.')

    doc.add_paragraph('\n// Use this block for each question when adding Multiple Choice Questions (MCQ)')

    # Question 1
    doc.add_paragraph('\n@question Each student at Riverside Prep chooses a uniform consisting of 1 shirt, 1 pair of pants, and 1 hat. The table shows the available options for each item. How many different complete uniforms are possible?')
    doc.add_paragraph('@instruction Select the number of possible uniform combinations from the options.')
    doc.add_paragraph('@difficulty easy')
    doc.add_paragraph('@Order 1')
    doc.add_paragraph('@option (A) 18')
    doc.add_paragraph('@option (B) 24')
    doc.add_paragraph('@@option (C) 36')
    doc.add_paragraph('@option (D) 48')
    doc.add_paragraph('@option (E) 72')
    doc.add_paragraph('@explanation ')
    doc.add_paragraph('There are 4 shirt choices, 3 pants choices, and 3 hat choices. Total combinations = 4 × 3 × 3 = 36.')
    doc.add_paragraph('@subject Quantitative Math')
    doc.add_paragraph('@unit Numbers and Operations')
    doc.add_paragraph('@topic Combinations / Counting')
    doc.add_paragraph('@plusmarks 1')
    doc.add_picture(str(img1), width=Inches(6))

    # Question 2
    doc.add_paragraph('\n@question The top view of a rectangular box containing 8 identical tightly packed spherical balls arranged in two rows of four is shown. If each ball has radius $2$ cm, which of the following is closest to the dimensions (height × width × length), in centimeters, of the rectangular package?')
    doc.add_paragraph('@instruction Choose the correct dimensions from the options.')
    doc.add_paragraph('@difficulty moderate')
    doc.add_paragraph('@Order 2')
    doc.add_paragraph('@option (A) $4 \\times 8 \\times 16$')
    doc.add_paragraph('@option (B) $2 \\times 8 \\times 16$')
    doc.add_paragraph('@@option (C) $4 \\times 12 \\times 18$')
    doc.add_paragraph('@option (D) $6 \\times 8 \\times 16$')
    doc.add_paragraph('@option (E) $8 \\times 12 \\times 24$')
    doc.add_paragraph('@explanation ')
    doc.add_paragraph('Top view shows two rows and four columns of circles (diameter = 4 cm). Width (short side) = 2 rows × 4 cm = 8 cm. Length = 4 columns × 4 cm = 16 cm. Height must accommodate one layer of balls: diameter = 4 cm. So dimensions: $4 \\times 8 \\times 16$.')
    doc.add_paragraph('@subject Quantitative Math')
    doc.add_paragraph('@unit Geometry and Measurement')
    doc.add_paragraph('@topic Solid Figures / Coordinate Geometry')
    doc.add_paragraph('@plusmarks 1')
    doc.add_picture(str(img2), width=Inches(6))

    doc.save(str(docx_path))
    print(f"Created {docx_path}")

def make_github_folder(base_dir: Path, img1: Path, img2: Path, docx_path: Path):
    repo_dir = base_dir / "github_repo"
    if repo_dir.exists():
        shutil.rmtree(repo_dir)
    repo_dir.mkdir(parents=True)

    images_dir = repo_dir / "images"
    images_dir.mkdir()

    shutil.copy(img1, images_dir / img1.name)
    shutil.copy(img2, images_dir / img2.name)

    readme = repo_dir / "README.md"
    readme.write_text("# Generated Assessment Questions\n\nThis repository contains two generated quantitative math questions in the requested Question Output Format, plus images.\n")
    questions_md = repo_dir / "QUESTIONS.md"
    questions_md.write_text(
        (
            "@title Central Middle — Derived Assessment\n"
            "@description Two generated quantitative math questions similar to the base examples.\n\n"
            "// Question 1\n"
            "@question Each student at Riverside Prep chooses a uniform consisting of 1 shirt, 1 pair of pants, and 1 hat. The table shows the available options for each item. How many different complete uniforms are possible?\n"
            "@instruction Select the number of possible uniform combinations from the options.\n"
            "@difficulty easy\n"
            "@Order 1\n"
            "@option (A) 18\n"
            "@option (B) 24\n"
            "@@option (C) 36\n"
            "@option (D) 48\n"
            "@option (E) 72\n"
            "@explanation \n"
            "There are 4 shirt choices, 3 pants choices, and 3 hat choices. Total combinations = 4 × 3 × 3 = 36.\n"
            "@subject Quantitative Math\n"
            "@unit Numbers and Operations\n"
            "@topic Combinations / Counting\n"
            "@plusmarks 1\n\n"
            "// Question 2\n"
            "@question The top view of a rectangular box containing 8 identical tightly packed spherical balls arranged in two rows of four is shown. If each ball has radius $2$ cm, which of the following is closest to the dimensions (height × width × length), in centimeters, of the rectangular package?\n"
            "@instruction Choose the correct dimensions from the options.\n"
            "@difficulty moderate\n"
            "@Order 2\n"
            "@option (A) $4 \\times 8 \\times 16$\n"
            "@option (B) $2 \\times 8 \\times 16$\n"
            "@@option (C) $4 \\times 12 \\times 18$\n"
            "@option (D) $6 \\times 8 \\times 16$\n"
            "@option (E) $8 \\times 12 \\times 24$\n"
            "@explanation \n"
            "Top view shows two rows and four columns of circles (diameter = 4 cm). Width (short side) = 2 rows × 4 cm = 8 cm. Length = 4 columns × 4 cm = 16 cm. Height must accommodate one layer of balls: diameter = 4 cm. So dimensions: $4 \\times 8 \\times 16$.\n"
            "@subject Quantitative Math\n"
            "@unit Geometry and Measurement\n"
            "@topic Solid Figures / Coordinate Geometry\n"
            "@plusmarks 1\n"
        )
    )

    # Copy docx as well
    shutil.copy(docx_path, repo_dir / docx_path.name)

    # Create zip
    zip_path = base_dir / "github_repo.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(repo_dir):
            for f in files:
                full = Path(root) / f
                zf.write(full, arcname=str(full.relative_to(repo_dir)))
    print(f"Created {zip_path}")
    return repo_dir, zip_path

def git_init_and_push(repo_dir: Path, remote_url: str, branch: str = "main"):
    """
    Initialize git repo in repo_dir (if not already initialized), commit and push to remote_url.
    remote_url must be a valid git remote that you have push access to.
    """
    cwd = str(repo_dir)
    def run(cmd):
        print("> " + " ".join(cmd))
        subprocess.run(cmd, cwd=cwd, check=True)

    # Init if needed
    if not (repo_dir / ".git").exists():
        run(["git", "init"])
    run(["git", "checkout", "-B", branch])
    run(["git", "add", "."])
    run(["git", "commit", "-m", "Add generated assessment questions"])
    run(["git", "remote", "remove", "origin"],) if subprocess.run(["git", "remote"], cwd=cwd, capture_output=True).returncode == 0 else None
    # set remote
    run(["git", "remote", "add", "origin", remote_url])
    # push
    run(["git", "push", "-u", "origin", branch])
    print("Pushed to remote:", remote_url)

def main():
    parser = argparse.ArgumentParser(description="Generate assessment doc + github-ready repo")
    parser.add_argument("--outdir", default="assessment_output", help="Output directory")
    parser.add_argument("--push-repo", default=None, help="Remote git repo URL to push (optional)")
    parser.add_argument("--branch", default="main", help="Branch name for push")
    args = parser.parse_args()

    base = Path(args.outdir).absolute()
    base.mkdir(parents=True, exist_ok=True)

    img1 = base / "uniform_table.png"
    img2 = base / "rect_package_topview_8.png"
    docx_path = base / "generated_assessment.docx"

    # Create images
    make_uniform_image(img1)
    make_packed_spheres_image(img2, rows=2, cols=4, radius_cm=2)

    # Create docx
    create_docx(docx_path, img1, img2)

    # Make github-ready folder + zip
    repo_dir, zip_path = make_github_folder(base, img1, img2, docx_path)

    print("All done. Files created in:", base)
    print(" - Docx:", docx_path)
    print(" - Repo folder:", repo_dir)
    print(" - Zip:", zip_path)

    if args.push_repo:
        # Confirm remote URL looks like git URL
        remote = args.push_repo
        try:
            git_init_and_push(repo_dir, remote, branch=args.branch)
        except subprocess.CalledProcessError as e:
            print("Git command failed. Ensure git is installed and you have network & push permissions.")
            raise

if __name__ == "__main__":
    main()
